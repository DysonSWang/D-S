#!/usr/bin/env python3
"""生成设备绑定功能需求文档 Word 版本"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('设备绑定功能需求文档 (PRD)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 项目信息
    doc.add_paragraph('项目名称：三诺医疗设备蓝牙绑定功能')
    doc.add_paragraph('版本：v1.1')
    doc.add_paragraph('最后更新：2026-03-03')
    doc.add_paragraph('文档状态：待研发')
    doc.add_paragraph()
    
    # 一、项目概述
    doc.add_heading('一、项目概述', level=1)
    
    doc.add_heading('1.1 背景', level=2)
    doc.add_paragraph('为用户提供便捷的蓝牙医疗设备（血糖仪、血压计）绑定功能，实现设备与 APP 的快速配对和数据同步。')
    
    doc.add_heading('1.2 目标设备', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['设备类型', '型号', '连接方式', '配套 APP']
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    devices = [
        ['血糖仪', '臻准 2000 (BA-2000)', '蓝牙 4.0 BLE', '三诺健康/诺安健康'],
        ['血压计', 'BA-803', '蓝牙 4.0 BLE', '三诺健康']
    ]
    for d in devices:
        row = table.add_row()
        for i, cell in enumerate(d):
            row.cells[i].text = cell
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 用户场景', level=2)
    scenarios = [
        '用户首次购买三诺设备，需要通过 APP 绑定设备',
        '用户更换手机后，需要重新绑定设备',
        '用户绑定后可自动同步测量数据到 APP'
    ]
    for s in scenarios:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_page_break()
    
    # 二、功能流程
    doc.add_heading('二、功能流程总览', level=1)
    
    flow_text = '选择设备类型 → 操作指引 → 搜索设备 → 选择设备 → 配对确认 → 绑定成功'
    doc.add_paragraph(flow_text)
    
    doc.add_paragraph('Step 1: 选择设备类型 → Step 2: 操作指引 → Step 3: 搜索设备 → Step 4: 选择设备 → Step 5: 配对确认 → Step 6: 绑定成功')
    
    doc.add_page_break()
    
    # 三、详细步骤
    doc.add_heading('三、详细步骤说明', level=1)
    
    # Step 1
    doc.add_heading('Step 1: 选择设备类型', level=2)
    doc.add_paragraph('【此处插入截图：step1_select_device.png】')
    doc.add_paragraph('截图说明：展示两个设备卡片（橙色血糖仪 + 青色血压计）')
    
    doc.add_heading('页面元素', level=3)
    elem_table = doc.add_table(rows=1, cols=3)
    elem_table.style = 'Table Grid'
    elem_headers = elem_table.rows[0].cells
    for i, h in enumerate(['元素', '描述', '样式']):
        elem_headers[i].text = h
        elem_headers[i].paragraphs[0].runs[0].bold = True
    
    elements = [
        ['导航栏', '返回按钮 + 页面标题"选择设备"', '白色背景，固定顶部'],
        ['进度条', '6 个点，第 1 个高亮', '青色高亮，灰色未激活'],
        ['血糖仪卡片', '橙色主题，显示设备图标和名称', '#FB9851 橙色渐变'],
        ['血压计卡片', '青色主题，显示设备图标和名称', '#00C2CC 青色渐变']
    ]
    for e in elements:
        row = elem_table.add_row()
        for i, cell in enumerate(e):
            row.cells[i].text = cell
    
    doc.add_paragraph()
    doc.add_heading('交互逻辑', level=3)
    doc.add_paragraph('• 点击设备卡片进入下一步')
    doc.add_paragraph('• 顶部导航栏支持返回')
    doc.add_paragraph('• 卡片有按压缩放效果 (scale: 0.98)')
    
    doc.add_page_break()
    
    # Step 2
    doc.add_heading('Step 2: 操作指引', level=2)
    doc.add_paragraph('【此处插入截图：step2_guide.png】')
    doc.add_paragraph('截图说明：展示 4 步图文指引和底部"已阅读，开始绑定"按钮')
    
    doc.add_heading('血糖仪指引内容', level=3)
    guides = [
        '步骤 1: 安装试纸或长按开机键 3 秒开机',
        '步骤 2: 确认屏幕显示蓝牙图标（闪烁表示可配对）',
        '步骤 3: 部分型号需进入"设置"→"蓝牙"菜单',
        '步骤 4: 保持设备与手机距离 ≤ 1 米'
    ]
    for g in guides:
        doc.add_paragraph(g, style='List Number')
    
    doc.add_heading('血压计指引内容', level=3)
    bp_guides = [
        '步骤 1: 安装电池 (4 节 7 号电池) 或连接 USB 供电',
        '步骤 2: 正确缠绕袖带',
        '步骤 3: 长按电源键 3 秒开机，确认蓝牙图标显示',
        '步骤 4: 部分型号需长按"记忆"键 5 秒进入配对模式'
    ]
    for g in bp_guides:
        doc.add_paragraph(g, style='List Number')
    
    doc.add_page_break()
    
    # Step 3
    doc.add_heading('Step 3: 搜索设备', level=2)
    doc.add_paragraph('【此处插入截图：step3_search.png】')
    doc.add_paragraph('截图说明：展示扩散圆环动画和"正在搜索设备..."文字')
    
    doc.add_heading('异常处理', level=3)
    exception_table = doc.add_table(rows=1, cols=2)
    exception_table.style = 'Table Grid'
    exc_headers = exception_table.rows[0].cells
    for i, h in enumerate(['异常场景', '处理方式']):
        exc_headers[i].text = h
        exc_headers[i].paragraphs[0].runs[0].bold = True
    
    exceptions = [
        ['蓝牙未开启', '弹窗提示用户开启蓝牙'],
        ['30 秒未找到设备', '显示"未找到设备"，提供重试按钮'],
        ['权限被拒绝', '引导用户到系统设置开启蓝牙权限']
    ]
    for e in exceptions:
        row = exception_table.add_row()
        for i, cell in enumerate(e):
            row.cells[i].text = cell
    
    doc.add_page_break()
    
    # Step 4
    doc.add_heading('Step 4: 选择设备', level=2)
    doc.add_paragraph('【此处插入截图：step4_select.png】')
    doc.add_paragraph('截图说明：展示设备列表，包含设备名称和 MAC 地址')
    
    doc.add_heading('设备名称格式', level=3)
    doc.add_paragraph('• 血糖仪：Sinocare_BA2000_XXXX 或 SN-XXXX')
    doc.add_paragraph('• 血压计：Sinocare_BA803_XXXX 或 BP-XXXX')
    doc.add_paragraph('• XXXX = MAC 地址后 4 位')
    
    doc.add_page_break()
    
    # Step 5
    doc.add_heading('Step 5: 配对确认', level=2)
    doc.add_paragraph('【此处插入截图：step5_pairing.png】')
    doc.add_paragraph('截图说明：展示配对码输入框或动态配对码显示')
    
    doc.add_heading('配对模式', level=3)
    pairing_table = doc.add_table(rows=1, cols=3)
    pairing_table.style = 'Table Grid'
    pair_headers = pairing_table.rows[0].cells
    for i, h in enumerate(['配对模式', '说明', '界面表现']):
        pair_headers[i].text = h
        pair_headers[i].paragraphs[0].runs[0].bold = True
    
    pairings = [
        ['无需配对码', '部分新型号设备', '直接显示"正在配对"，自动完成'],
        ['固定配对码', '老款设备', '提示用户输入 000000 或 123456'],
        ['动态配对码', '部分型号', '显示设备屏幕上的 6 位数字']
    ]
    for p in pairings:
        row = pairing_table.add_row()
        for i, cell in enumerate(p):
            row.cells[i].text = cell
    
    doc.add_page_break()
    
    # Step 6
    doc.add_heading('Step 6: 绑定成功', level=2)
    doc.add_paragraph('【此处插入截图：step6_success.png】')
    doc.add_paragraph('截图说明：展示成功图标、设备信息卡片和"完成"按钮')
    
    doc.add_page_break()
    
    # 四、技术规格
    doc.add_heading('四、技术规格', level=1)
    
    doc.add_heading('4.1 蓝牙技术要求', level=2)
    tech_specs = [
        '协议：Bluetooth 4.0 BLE (低功耗蓝牙)',
        '广播间隔：20ms - 500ms (设备端配置)',
        '配对距离：≤ 1 米 (建议)',
        '连接超时：30 秒 (搜索) / 60 秒 (配对)'
    ]
    for spec in tech_specs:
        doc.add_paragraph(spec, style='List Bullet')
    
    doc.add_heading('4.2 状态码定义', level=2)
    status_table = doc.add_table(rows=1, cols=2)
    status_table.style = 'Table Grid'
    status_headers = status_table.rows[0].cells
    for i, h in enumerate(['状态码', '说明']):
        status_headers[i].text = h
        status_headers[i].paragraphs[0].runs[0].bold = True
    
    statuses = [
        ['SUCCESS', '操作成功'],
        ['BLUETOOTH_OFF', '蓝牙未开启'],
        ['PERMISSION_DENIED', '权限被拒绝'],
        ['SEARCH_TIMEOUT', '搜索超时'],
        ['PAIRING_TIMEOUT', '配对超时'],
        ['PAIRING_CODE_ERROR', '配对码错误'],
        ['CONNECTION_FAILED', '连接失败'],
        ['DEVICE_NOT_FOUND', '设备未找到']
    ]
    for s in statuses:
        row = status_table.add_row()
        for i, cell in enumerate(s):
            row.cells[i].text = cell
    
    doc.add_page_break()
    
    # 五、UI/UX 规范
    doc.add_heading('五、UI/UX 规范', level=1)
    
    doc.add_heading('5.1 颜色规范', level=2)
    color_table = doc.add_table(rows=1, cols=3)
    color_table.style = 'Table Grid'
    color_headers = color_table.rows[0].cells
    for i, h in enumerate(['颜色用途', '色值', '说明']):
        color_headers[i].text = h
        color_headers[i].paragraphs[0].runs[0].bold = True
    
    colors = [
        ['主色调', '#00C2CC', '青色，用于血压计主题'],
        ['血糖仪主题色', '#FB9851', '橙色'],
        ['血压计主题色', '#00C2CC', '青色'],
        ['成功色', '#00C2CC', '与主色调一致'],
        ['错误色', '#F04838', '红色'],
        ['警告色', '#FB7A1E', '橙色']
    ]
    for c in colors:
        row = color_table.add_row()
        for i, cell in enumerate(c):
            row.cells[i].text = cell
    
    doc.add_paragraph()
    
    # 六、测试要求
    doc.add_heading('六、测试要求', level=1)
    
    doc.add_heading('6.1 功能测试', level=2)
    func_tests = [
        '血糖仪完整绑定流程',
        '血压计完整绑定流程',
        '搜索超时处理',
        '配对码验证 (固定/动态/无需)',
        '绑定成功后设备列表展示',
        '已绑定设备自动重连'
    ]
    for t in func_tests:
        doc.add_paragraph(f'□ {t}', style='List Bullet')
    
    doc.add_heading('6.2 兼容性测试', level=2)
    compat_tests = [
        'iOS 12+ 系统',
        'Android 8+ 系统',
        '不同屏幕尺寸适配',
        '深色模式适配 (可选)'
    ]
    for t in compat_tests:
        doc.add_paragraph(f'□ {t}', style='List Bullet')
    
    doc.add_page_break()
    
    # 七、交付物
    doc.add_heading('七、交付物', level=1)
    
    doc.add_heading('7.1 前端交付', level=2)
    deliverables = [
        '完整 HTML/CSS/JS 原型 (参考 device_binding_demo.html)',
        '6 个页面的完整交互逻辑',
        '响应式适配 (移动端)',
        '加载动画和过渡效果'
    ]
    for d in deliverables:
        doc.add_paragraph(f'□ {d}', style='List Bullet')
    
    doc.add_heading('7.2 接口文档', level=2)
    api_docs = [
        '蓝牙扫描 API 文档',
        '蓝牙配对 API 文档',
        '设备绑定保存 API 文档',
        '错误码定义文档'
    ]
    for d in api_docs:
        doc.add_paragraph(f'□ {d}', style='List Bullet')
    
    doc.add_heading('7.3 测试报告', level=2)
    test_reports = [
        '功能测试报告',
        '兼容性测试报告',
        '设备兼容性测试报告'
    ]
    for d in test_reports:
        doc.add_paragraph(f'□ {d}', style='List Bullet')
    
    doc.add_page_break()
    
    # 八、附录
    doc.add_heading('八、附录', level=1)
    
    doc.add_heading('8.1 截图指南', level=2)
    doc.add_paragraph('1. 在浏览器中打开 device_binding_demo.html')
    doc.add_paragraph('2. 按步骤截图每个页面')
    doc.add_paragraph('3. 将截图保存到 screenshots/ 目录')
    doc.add_paragraph('4. 在 Word 文档中插入对应截图')
    
    doc.add_heading('建议的截图尺寸', level=3)
    doc.add_paragraph('• 移动端尺寸：375x812 (iPhone X) 或 360x640 (Android)')
    doc.add_paragraph('• 格式：PNG (支持透明背景)')
    doc.add_paragraph('• 命名：step{N}_{description}.png')
    
    doc.add_heading('8.2 术语表', level=2)
    term_table = doc.add_table(rows=1, cols=2)
    term_table.style = 'Table Grid'
    term_headers = term_table.rows[0].cells
    for i, h in enumerate(['术语', '说明']):
        term_headers[i].text = h
        term_headers[i].paragraphs[0].runs[0].bold = True
    
    terms = [
        ['BLE', 'Bluetooth Low Energy，低功耗蓝牙'],
        ['广播', '设备发送广播包，让手机可以发现'],
        ['配对', '建立安全连接的过程'],
        ['绑定', '保存配对信息，实现自动重连'],
        ['RSSI', '接收信号强度指示']
    ]
    for t in terms:
        row = term_table.add_row()
        for i, cell in enumerate(t):
            row.cells[i].text = cell
    
    # 保存文档
    output_path = '/root/.openclaw/workspace/设备绑定功能需求文档.docx'
    doc.save(output_path)
    print(f'Word 文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
