#!/usr/bin/env python3
"""
生成完整的带截图 Word 文档
"""

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_complete_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('设备绑定功能需求文档 (PRD)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 项目信息
    doc.add_paragraph('项目名称：三诺医疗设备蓝牙绑定功能')
    doc.add_paragraph('版本：v1.1')
    doc.add_paragraph('最后更新：2026-03-03')
    doc.add_paragraph('文档状态：待研发')
    doc.add_paragraph()
    
    # 一、项目概述
    doc.add_heading('一、项目概述', level=1)
    
    doc.add_heading('1.1 背景', level=2)
    doc.add_paragraph('为用户提供便捷的蓝牙医疗设备（血糖仪、血压计）绑定功能，实现设备与 APP 的快速配对和数据同步。')
    
    doc.add_heading('1.2 目标设备', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['设备类型', '型号', '连接方式', '配套 APP']
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    devices = [
        ['血糖仪', '臻准 2000 (BA-2000)', '蓝牙 4.0 BLE', '三诺健康/诺安健康'],
        ['血压计', 'BA-803', '蓝牙 4.0 BLE', '三诺健康']
    ]
    for d in devices:
        row = table.add_row()
        for i, cell in enumerate(d):
            row.cells[i].text = cell
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 用户场景', level=2)
    scenarios = [
        '用户首次购买三诺设备，需要通过 APP 绑定设备',
        '用户更换手机后，需要重新绑定设备',
        '用户绑定后可自动同步测量数据到 APP'
    ]
    for s in scenarios:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_page_break()
    
    # 二、功能流程
    doc.add_heading('二、功能流程总览', level=1)
    doc.add_paragraph('选择设备类型 → 操作指引 → 搜索设备 → 选择设备 → 配对确认 → 绑定成功')
    doc.add_paragraph()
    
    # 三、详细步骤
    doc.add_heading('三、详细步骤说明', level=1)
    
    screenshot_dir = '/root/.openclaw/workspace/screenshots'
    
    # Step 1
    doc.add_heading('Step 1: 选择设备类型', level=2)
    img_path = os.path.join(screenshot_dir, 'step1_select_device.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('说明：用户选择要绑定的设备类型（血糖仪或血压计）')
    doc.add_page_break()
    
    # Step 2
    doc.add_heading('Step 2: 操作指引', level=2)
    img_path = os.path.join(screenshot_dir, 'step2_guide_glucose.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('说明：血糖仪操作指引（血压计指引类似）')
    doc.add_paragraph('指引内容包括：')
    doc.add_paragraph('• 安装试纸或长按开机键 3 秒开机', style='List Bullet')
    doc.add_paragraph('• 确认屏幕显示蓝牙图标（闪烁表示可配对）', style='List Bullet')
    doc.add_paragraph('• 部分型号需进入"设置"→"蓝牙"菜单', style='List Bullet')
    doc.add_paragraph('• 保持设备与手机距离 ≤ 1 米', style='List Bullet')
    doc.add_page_break()
    
    # Step 3
    doc.add_heading('Step 3: 搜索设备', level=2)
    img_path = os.path.join(screenshot_dir, 'step3_searching.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('说明：显示扩散圆环动画，自动搜索附近设备')
    doc.add_paragraph('超时时间：30 秒')
    doc.add_page_break()
    
    # Step 4
    doc.add_heading('Step 4: 选择设备', level=2)
    img_path = os.path.join(screenshot_dir, 'step4_device_list.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('说明：显示搜索到的设备列表，用户选择要连接的设备')
    doc.add_paragraph('设备名称格式：Sinocare_BA2000_XXXX (XXXX 为 MAC 地址后 4 位)')
    doc.add_page_break()
    
    # Step 5
    doc.add_heading('Step 5: 配对确认', level=2)
    img_path = os.path.join(screenshot_dir, 'step5_pairing.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('说明：显示配对码，用户确认设备屏幕上的数字一致')
    doc.add_paragraph('配对模式：')
    doc.add_paragraph('• 无需配对码：部分新型号设备', style='List Bullet')
    doc.add_paragraph('• 固定配对码：老款设备 (000000 或 123456)', style='List Bullet')
    doc.add_paragraph('• 动态配对码：部分型号 (6 位数字)', style='List Bullet')
    doc.add_page_break()
    
    # Step 6
    doc.add_heading('Step 6: 绑定成功', level=2)
    img_path = os.path.join(screenshot_dir, 'step6_success.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(14))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('说明：显示绑定成功状态和设备信息')
    doc.add_paragraph('绑定信息保存：')
    doc.add_paragraph('• 设备 MAC 地址', style='List Bullet')
    doc.add_paragraph('• 设备型号', style='List Bullet')
    doc.add_paragraph('• 绑定时间', style='List Bullet')
    
    doc.add_page_break()
    
    # 四、技术规格
    doc.add_heading('四、技术规格', level=1)
    
    doc.add_heading('4.1 蓝牙技术要求', level=2)
    tech_specs = [
        '协议：Bluetooth 4.0 BLE (低功耗蓝牙)',
        '广播间隔：20ms - 500ms (设备端配置)',
        '配对距离：≤ 1 米 (建议)',
        '连接超时：30 秒 (搜索) / 60 秒 (配对)'
    ]
    for spec in tech_specs:
        doc.add_paragraph(spec, style='List Bullet')
    
    doc.add_heading('4.2 状态码定义', level=2)
    status_table = doc.add_table(rows=1, cols=2)
    status_table.style = 'Table Grid'
    status_headers = status_table.rows[0].cells
    for i, h in enumerate(['状态码', '说明']):
        status_headers[i].text = h
        status_headers[i].paragraphs[0].runs[0].bold = True
    
    statuses = [
        ['SUCCESS', '操作成功'],
        ['BLUETOOTH_OFF', '蓝牙未开启'],
        ['PERMISSION_DENIED', '权限被拒绝'],
        ['SEARCH_TIMEOUT', '搜索超时'],
        ['PAIRING_TIMEOUT', '配对超时'],
        ['PAIRING_CODE_ERROR', '配对码错误'],
        ['CONNECTION_FAILED', '连接失败'],
        ['DEVICE_NOT_FOUND', '设备未找到']
    ]
    for s in statuses:
        row = status_table.add_row()
        for i, cell in enumerate(s):
            row.cells[i].text = cell
    
    doc.add_page_break()
    
    # 五、UI/UX 规范
    doc.add_heading('五、UI/UX 规范', level=1)
    
    doc.add_heading('5.1 颜色规范', level=2)
    color_table = doc.add_table(rows=1, cols=3)
    color_table.style = 'Table Grid'
    color_headers = color_table.rows[0].cells
    for i, h in enumerate(['颜色用途', '色值', '说明']):
        color_headers[i].text = h
        color_headers[i].paragraphs[0].runs[0].bold = True
    
    colors = [
        ['主色调', '#00C2CC', '青色，用于血压计主题'],
        ['血糖仪主题色', '#FB9851', '橙色'],
        ['血压计主题色', '#00C2CC', '青色'],
        ['成功色', '#00C2CC', '与主色调一致'],
        ['错误色', '#F04838', '红色'],
        ['警告色', '#FB7A1E', '橙色']
    ]
    for c in colors:
        row = color_table.add_row()
        for i, cell in enumerate(c):
            row.cells[i].text = cell
    
    doc.add_paragraph()
    
    # 六、测试要求
    doc.add_heading('六、测试要求', level=1)
    
    doc.add_heading('6.1 功能测试', level=2)
    func_tests = [
        '血糖仪完整绑定流程',
        '血压计完整绑定流程',
        '搜索超时处理',
        '配对码验证 (固定/动态/无需)',
        '绑定成功后设备列表展示',
        '已绑定设备自动重连'
    ]
    for t in func_tests:
        doc.add_paragraph(f'□ {t}', style='List Bullet')
    
    doc.add_heading('6.2 兼容性测试', level=2)
    compat_tests = [
        'iOS 12+ 系统',
        'Android 8+ 系统',
        '不同屏幕尺寸适配',
        '深色模式适配 (可选)'
    ]
    for t in compat_tests:
        doc.add_paragraph(f'□ {t}', style='List Bullet')
    
    doc.add_page_break()
    
    # 七、交付物
    doc.add_heading('七、交付物', level=1)
    
    doc.add_heading('7.1 前端交付', level=2)
    deliverables = [
        '完整 HTML/CSS/JS 原型 (参考 device_binding_demo.html)',
        '6 个页面的完整交互逻辑',
        '响应式适配 (移动端)',
        '加载动画和过渡效果'
    ]
    for d in deliverables:
        doc.add_paragraph(f'□ {d}', style='List Bullet')
    
    doc.add_heading('7.2 接口文档', level=2)
    api_docs = [
        '蓝牙扫描 API 文档',
        '蓝牙配对 API 文档',
        '设备绑定保存 API 文档',
        '错误码定义文档'
    ]
    for d in api_docs:
        doc.add_paragraph(f'□ {d}', style='List Bullet')
    
    doc.add_heading('7.3 测试报告', level=2)
    test_reports = [
        '功能测试报告',
        '兼容性测试报告',
        '设备兼容性测试报告'
    ]
    for d in test_reports:
        doc.add_paragraph(f'□ {d}', style='List Bullet')
    
    doc.add_page_break()
    
    # 八、附录
    doc.add_heading('八、附录', level=1)
    
    doc.add_heading('8.1 术语表', level=2)
    term_table = doc.add_table(rows=1, cols=2)
    term_table.style = 'Table Grid'
    term_headers = term_table.rows[0].cells
    for i, h in enumerate(['术语', '说明']):
        term_headers[i].text = h
        term_headers[i].paragraphs[0].runs[0].bold = True
    
    terms = [
        ['BLE', 'Bluetooth Low Energy，低功耗蓝牙'],
        ['广播', '设备发送广播包，让手机可以发现'],
        ['配对', '建立安全连接的过程'],
        ['绑定', '保存配对信息，实现自动重连'],
        ['RSSI', '接收信号强度指示']
    ]
    for t in terms:
        row = term_table.add_row()
        for i, cell in enumerate(t):
            row.cells[i].text = cell
    
    doc.add_paragraph()
    doc.add_heading('8.2 参考资料', level=2)
    doc.add_paragraph('• 调研报告：device_binding_research.md')
    doc.add_paragraph('• HTML 原型：device_binding_demo.html')
    doc.add_paragraph('• 三诺官网：https://www.sinocare.com')
    doc.add_paragraph('• 蓝牙规范：Bluetooth Core Specification v5.0')
    
    # 保存文档
    output_path = '/root/.openclaw/workspace/设备绑定功能需求文档 - 完整版.docx'
    doc.save(output_path)
    print(f'Word 文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_complete_document()
