#!/usr/bin/env python3
"""
更新 Word 文档，插入真实截图
"""

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def update_document():
    doc_path = '/root/.openclaw/workspace/设备绑定功能需求文档.docx'
    screenshot_dir = '/root/.openclaw/workspace/screenshots'
    output_path = '/root/.openclaw/workspace/设备绑定功能需求文档 - 完整版.docx'
    
    doc = Document(doc_path)
    
    # 截图映射
    screenshots = {
        'Step 1': 'step1_select_device.png',
        'Step 2': 'step2_guide_glucose.png',
        'Step 3': 'step3_searching.png',
        'Step 4': 'step4_device_list.png',
        'Step 5': 'step5_pairing.png',
        'Step 6': 'step6_success.png',
    }
    
    print(f'正在处理文档：{doc_path}')
    
    # 遍历所有段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 查找占位符
        for step_name, screenshot_file in screenshots.items():
            if f'【此处插入截图' in text and screenshot_file.split('.')[0] in text:
                print(f'找到 {step_name} 占位符，准备插入 {screenshot_file}')
                
                screenshot_path = os.path.join(screenshot_dir, screenshot_file)
                if os.path.exists(screenshot_path):
                    # 在占位符前插入图片
                    run = para.add_run()
                    run.add_picture(screenshot_path, width=Cm(14))
                    print(f'✓ 已插入：{screenshot_file}')
                else:
                    print(f'✗ 截图文件不存在：{screenshot_path}')
                break
    
    # 保存新文档
    doc.save(output_path)
    print(f'\n完成！文档已保存到：{output_path}')
    return output_path

if __name__ == '__main__':
    update_document()
